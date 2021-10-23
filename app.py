import pandas as pd
import numpy as np
import streamlit as st
#######################
from selenium import webdriver 
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from getpass import getpass
import time
#from Screenshot import Screenshot_Clipping
########################
from docx import Document
from docx.enum.section import WD_ORIENT
import docx
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from datetime import date
#########################
from PIL import Image
import PIL
from io import BytesIO
#########################
import docx2txt
import textract
#########################
import base64
from io import BytesIO
#########################
import sqlite3
import pandas as pd
#########################
import hashlib
import re
#########################
import zipfile, os
#########################

### Hide module on App ----------------
hide_menu = """
<style>
#MainMenu {
    visibility:visible;
}

footer{
    visibility:hidden;
}

</style>

"""






####### ------------------ SECURITY -------------------------------- #######

            ######### ---- HASHING ------- ########

def make_hashes(password):
    return hashlib.sha256(str.encode(password)).hexdigest()

def check_hashes(password, hashed_text):
    if make_hashes(password) == hashed_text:
        return hashed_text
    return False



############################################################

## ---------------- DB MANAGER ----------------------- ###
conn = sqlite3.connect('data.db4')
c = conn.cursor()

conn = sqlite3.connect('data.db4')
c = conn.cursor()
def create_usertable():
	c.execute('CREATE TABLE IF NOT EXISTS userstable(username TEXT,password TEXT)')


def add_userdata(username,password):
	c.execute('INSERT INTO userstable(username,password) VALUES (?,?)',(username,password))
	conn.commit()

def login_user(username,password):
	c.execute('SELECT * FROM userstable WHERE username =? AND password = ?',(username,password))
	data = c.fetchall()
	return data

def view_all_users():
    c.execute('SELECT * FROM userstable')
    data = c.fetchall()
    return data

############################################################


### -------------- CROPING FUNCTIONS --------------------

#### ---------------- GOLD_CodirWorldwide --------------------- #### #1
def Crop_GOLD_CodirWorldwide(): 
    # ---------------- Dash1
    with Image.open("worldwide_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_worldwide_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("worldwide_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_worldwide_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_worldwide_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_worldwide_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_worldwide_dash2_bleu.png', quality=95) # BLEU
    
    # ---------------- Dash3
    with Image.open("worldwide_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_worldwide_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_worldwide_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_worldwide_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_worldwide_dash3_bleu.png', quality=95) # BLEU


#### ---------------- GOLD_ZoneNorth Asia --------------------- #### #2
def Crop_ZoneNorth_Asia(): 
    # ---------------- Dash1
    with Image.open("ZoneNorth_Asia_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_ZoneNorth_Asia_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("ZoneNorth_Asia_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_ZoneNorth_Asia_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_ZoneNorth_Asia_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_ZoneNorth_Asia_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_ZoneNorth_Asia_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_ZoneNorth_America --------------------- #### #3
def Crop_ZoneNorth_America(): 
    # ---------------- Dash1
    with Image.open("ZoneNorth_America_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_ZoneNorth_America_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("ZoneNorth_America_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_ZoneNorth_America_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_ZoneNorth_America_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_ZoneNorth_America_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_ZoneNorth_America_dash2_bleu.png', quality=95) # BLEU
        
#### ----------------  GOLD_ZoneTravel_Retail --------------------- #### #4
def Crop_GOLD_ZoneTravel_Retail(): 
    # ---------------- Dash1
    with Image.open("ZoneTravel_Retail_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_ZoneTravel_Retail_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("ZoneTravel_Retail_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_ZoneTravel_Retail_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_ZoneTravel_Retail_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_ZoneTravel_Retail_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_ZoneTravel_Retail_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_ZoneWestern_Europe --------------------- #### #5
def Crop_GOLD_ZoneWestern_Europe(): 
    # ---------------- Dash1
    with Image.open("ZoneWestern_Europe_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_ZoneWestern_Europe_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("ZoneWestern_Europe_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_ZoneWestern_Europe_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_ZoneWestern_Europe_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_ZoneWestern_Europe_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_ZoneWestern_Europe_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_ZoneEastern_Europe --------------------- #### #6
def Crop_GOLD_ZoneEastern_Europe(): 
    # ---------------- Dash1
    with Image.open("ZoneEastern_Europe_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_ZoneEastern_Europe_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("ZoneEastern_Europe_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_ZoneEastern_Europe_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_ZoneEastern_Europe_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_ZoneEastern_Europe_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_ZoneEastern_Europe_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_ZoneSAPMENA --------------------- #### #7
def Crop_GOLD_ZoneSAPMENA(): 
    # ---------------- Dash1
    with Image.open("ZoneSAPMENA_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_ZoneSAPMENA_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("ZoneSAPMENA_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_ZoneSAPMENA_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_ZoneSAPMENA_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_ZoneSAPMENA_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_ZoneSAPMENA_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_ZoneLatin_America --------------------- #### #8
def Crop_GOLD_ZoneLatin_America(): 
    # ---------------- Dash1
    with Image.open("ZoneLatin_America_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_ZoneLatin_America_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("ZoneLatin_America_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_ZoneLatin_America_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_ZoneLatin_America_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_ZoneLatin_America_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_ZoneLatin_America_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_ZoneSSA --------------------- #### #9
def Crop_GOLD_ZoneSSA(): 
    # ---------------- Dash1
    with Image.open("ZoneSSA_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_ZoneSSA_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("ZoneSSA_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_ZoneSSA_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_ZoneSSA_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_ZoneSSA_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_ZoneSSA_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandKiehls --------------------- #### #10
def Crop_GOLD_BrandKiehls(): 
    # ---------------- Dash1
    with Image.open("BrandKiehls_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandKiehls_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandKiehls_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandKiehls_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandKiehls_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandKiehls_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandKiehls_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandKiehls_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandKiehls_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandKiehls_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandKiehls_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandKiehls_dash3_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandYves_Saint_Lauren --------------------- #### #11
def Crop_GOLD_BrandYves_Saint_Lauren(): 
    # ---------------- Dash1
    with Image.open("BrandYves_Saint_Laurent_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandYves_Saint_Laurent_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandYves_Saint_Laurent_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandYves_Saint_Laurent_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandYves_Saint_Laurent_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandYves_Saint_Laurent_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandYves_Saint_Laurent_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandYves_Saint_Laurent_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandYves_Saint_Laurent_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandYves_Saint_Laurent_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandYves_Saint_Laurent_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandYves_Saint_Laurent_dash3_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandArmani --------------------- #### #12
def Crop_GOLD_BrandArmani(): 
    # ---------------- Dash1
    with Image.open("BrandArmani_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandArmani_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandArmani_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandArmani_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandArmani_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandArmani_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandArmani_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandArmani_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandArmani_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandArmani_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandArmani_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandArmani_dash3_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandH_Rubinstein --------------------- #### #13
def Crop_GOLD_BrandH_Rubinstein(): 
    # ---------------- Dash1
    with Image.open("BrandH_Rubinstein_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandH_Rubinstein_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandH_Rubinstein_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandH_Rubinstein_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandH_Rubinstein_dash2_violet.png', quality=95) # VIOLET
        
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandH_Rubinstein_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandH_Rubinstein_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandH_Rubinstein_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandH_Rubinstein_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandH_Rubinstein_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandH_Rubinstein_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandH_Rubinstein_dash3_bleu.png', quality=95) # BLEU

#### ---------------- GOLD_BrandBiotherm --------------------- #### #14
def Crop_GOLD_BrandBiotherm(): 
    # ---------------- Dash1
    with Image.open("BrandBiotherm_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandBiotherm_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandBiotherm_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandBiotherm_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandBiotherm_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandBiotherm_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandBiotherm_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandBiotherm_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandBiotherm_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandBiotherm_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandBiotherm_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandBiotherm_dash3_bleu.png', quality=95) # BLEU

#### ---------------- GOLD_BrandIT_Cosmetics --------------------- #### #15
def Crop_GOLD_BrandIT_Cosmetics(): 
    # ---------------- Dash1
    with Image.open("BrandIT_Cosmetics_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandIT_Cosmetics_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandIT_Cosmetics_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandIT_Cosmetics_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandIT_Cosmetics_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandIT_Cosmetics_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandIT_Cosmetics_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandIT_Cosmetics_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandIT_Cosmetics_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandIT_Cosmetics_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandIT_Cosmetics_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandIT_Cosmetics_dash3_bleu.png', quality=95) # BLEU

#### ---------------- GOLD_BrandUrban_Decay --------------------- #### 16
def Crop_GOLD_BrandUrban_Decay(): 
    # ---------------- Dash1
    with Image.open("BrandUrban_Decay_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandUrban_Decay_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandUrban_Decay_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandUrban_Decay_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandUrban_Decay_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandUrban_Decay_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandUrban_Decay_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandUrban_Decay_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandUrban_Decay_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandUrban_Decay_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandUrban_Decay_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandUrban_Decay_dash3_bleu.png', quality=95) # BLEU

#### ---------------- GOLD_BrandShu_Uemura --------------------- #### 17
def Crop_GOLD_BrandShu_Uemura(): 
    # ---------------- Dash1
    with Image.open("BrandShu_Uemura_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandShu_Uemura_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandShu_Uemura_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandShu_Uemura_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandShu_Uemura_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandShu_Uemura_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandShu_Uemura_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandShu_Uemura_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandShu_Uemura_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandShu_Uemura_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandShu_Uemura_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandShu_Uemura_dash3_bleu.png', quality=95) # BLEU

#### ---------------- GOLD_BrandRalph_Lauren --------------------- #### 18
def Crop_GOLD_BrandRalph_Lauren(): 
    # ---------------- Dash1
    with Image.open("BrandRalph_Lauren_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandRalph_Lauren_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandRalph_Lauren_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandRalph_Lauren_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandRalph_Lauren_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandRalph_Lauren_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandRalph_Lauren_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandRalph_Lauren_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandRalph_Lauren_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandRalph_Lauren_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandRalph_Lauren_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandRalph_Lauren_dash3_bleu.png', quality=95) # BLEU

#### ---------------- GOLD_BrandYue_Sai --------------------- #### 19
def Crop_GOLD_BrandYue_Sai(): 
    # ---------------- Dash1
    with Image.open("BrandYue_Sai_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandYue_Sai_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandYue_Sai_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandYue_Sai_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandYue_Sai_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandYue_Sai_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandYue_Sai_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandYue_Sai_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandYue_Sai_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandYue_Sai_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandYue_Sai_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandYue_Sai_dash3_bleu.png', quality=95) # BLEU

#### ---------------- GOLD_BrandValentino --------------------- #### 20
def Crop_GOLD_BrandValentino(): 
    # ---------------- Dash1
    with Image.open("BrandValentino_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandValentino_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandValentino_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandValentino_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandValentino_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandValentino_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandValentino_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandValentino_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandValentino_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandValentino_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandValentino_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandValentino_dash3_bleu.png', quality=95) # BLEU


#### ---------------- GOLD_BrandViktor_and_Rolf --------------------- #### 21
def Crop_GOLD_BrandViktor_and_Rolf(): 
    # ---------------- Dash1
    with Image.open("BrandViktor_and_Rolf_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandViktor_and_Rolf_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandValentino_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandViktor_and_Rolf_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandViktor_and_Rolf_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandViktor_and_Rolf_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandViktor_and_Rolf_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandViktor_and_Rolf_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandViktor_and_Rolf_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandViktor_and_Rolf_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandViktor_and_Rolf_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandViktor_and_Rolf_dash3_bleu.png', quality=95) # BLEU

#### ---------------- GOLD_BrandAtelier_Cologne --------------------- #### 22
def Crop_GOLD_BrandAtelier_Cologne(): 
    # ---------------- Dash1
    with Image.open("BrandAtelier_Cologne_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandAtelier_Cologne_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandAtelier_Cologne_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandAtelier_Cologne_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandAtelier_Cologne_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandAtelier_Cologne_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandAtelier_Cologne_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandAtelier_Cologne_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandAtelier_Cologne_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandAtelier_Cologne_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandAtelier_Cologne_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandAtelier_Cologne_dash3_bleu.png', quality=95) # BLEU

#### ---------------- GOLD_BrandMaison_Margiela --------------------- #### 23
def Crop_GOLD_BrandMaison_Margiela(): 
    # ---------------- Dash1
    with Image.open("BrandMaison_Margiela_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandMaison_Margiela_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandMaison_Margiela_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandMaison_Margiela_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandMaison_Margiela_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandMaison_Margiela_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandMaison_Margiela_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandMaison_Margiela_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandMaison_Margiela_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandMaison_Margiela_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandMaison_Margiela_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandMaison_Margiela_dash3_bleu.png', quality=95) # BLEU

#### ---------------- GOLD_BrandMugler --------------------- #### 24
def Crop_GOLD_BrandMugler(): 
    # ---------------- Dash1
    with Image.open("BrandMugler_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandMugler_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandMugler_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandMugler_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandMugler_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandMugler_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandMugler_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandMugler_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandMugler_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandMugler_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandMugler_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandMugler_dash3_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandCacharel --------------------- #### 25
def Crop_GOLD_BrandCacharel(): 
    # ---------------- Dash1
    with Image.open("BrandCacharel_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandCacharel_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandCacharel_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandCacharel_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandCacharel_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandCacharel_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandCacharel_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandCacharel_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandCacharel_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandCacharel_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandCacharel_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandCacharel_dash3_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandDiesel --------------------- #### 26
def Crop_GOLD_BrandDiesel(): 
    # ---------------- Dash1
    with Image.open("BrandDiesel_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandDiesel_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandDiesel_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandDiesel_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandDiesel_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandDiesel_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandDiesel_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandDiesel_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandDiesel_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandDiesel_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandDiesel_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandDiesel_dash3_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandAzzaro --------------------- #### 27
def Crop_GOLD_BrandAzzaro(): 
    # ---------------- Dash1
    with Image.open("BrandAzzaro_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandAzzaro_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandAzzaro_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandAzzaro_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandAzzaro_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandAzzaro_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandAzzaro_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandAzzaro_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandAzzaro_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandAzzaro_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandAzzaro_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandAzzaro_dash3_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandPrada --------------------- #### 28
def Crop_GOLD_BrandPrada(): 
    # ---------------- Dash1
    with Image.open("BrandPrada_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandPrada_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandPrada_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandPrada_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandPrada_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandPrada_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandPrada_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandPrada_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandPrada_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandPrada_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandPrada_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandPrada_dash3_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandHouse_99 --------------------- #### 29
def Crop_GOLD_BrandHouse_99(): 
    # ---------------- Dash1
    with Image.open("BrandHouse_99_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandHouse_99_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandHouse_99_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandHouse_99_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandHouse_99_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandHouse_99_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandHouse_99_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandHouse_99_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandHouse_99_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandHouse_99_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandHouse_99_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandHouse_99_dash3_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandPaloma_Picasso --------------------- #### 30
def Crop_GOLD_BrandPaloma_Picasso(): 
    # ---------------- Dash1
    with Image.open("BrandPaloma_Picasso_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandPaloma_Picasso_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandPaloma_Picasso_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandPaloma_Picasso_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandPaloma_Picasso_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandPaloma_Picasso_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandPaloma_Picasso_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandPaloma_Picasso_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandPaloma_Picasso_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandPaloma_Picasso_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandPaloma_Picasso_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandPaloma_Picasso_dash3_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_BrandProenza_Schouler --------------------- #### 31
def Crop_GOLD_BrandProenza_Schouler(): 
    # ---------------- Dash1
    with Image.open("BrandProenza_Schouler_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_BrandProenza_Schouler_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("BrandProenza_Schouler_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_BrandProenza_Schouler_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_BrandProenza_Schouler_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_BrandProenza_Schouler_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_BrandProenza_Schouler_dash2_bleu.png', quality=95) # BLEU
        
     # ---------------- Dash3
    with Image.open("BrandProenza_Schouler_dash3.png") as im3: 
        im_Crop6 = im3.crop((0, 250, 1450, 1320))
        im_Crop6.save('crop_BrandProenza_Schouler_dash3_jaune.png', quality=95) # JAUNE
        
        im_Crop7 = im3.crop((0, 1320, 1450, 2020))
        im_Crop7.save('crop_BrandProenza_Schouler_dash3_violet.png', quality=95) # VIOLET
        
        im_Crop8 = im3.crop((0, 2020, 1450, 2770))
        im_Crop8.save('crop_BrandProenza_Schouler_dash3_vert.png', quality=95) # VERT
        
        im_Crop9 = im3.crop((0, 2770, 1450, 3910))
        im_Crop9.save('crop_BrandProenza_Schouler_dash3_bleu.png', quality=95) # BLEU


#### ----------------  GOLD_CountryChina --------------------- #### 32
def Crop_GOLD_CountryChina(): 
    # ---------------- Dash1
    with Image.open("CountryChina_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_CountryChina_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("CountryChina_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_CountryChina_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_CountryChina_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_CountryChina_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_CountryChina_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_CountryUnited_States --------------------- #### 33
def Crop_GOLD_CountryUnited_States(): 
    # ---------------- Dash1
    with Image.open("CountryUnited_States_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_CountryUnited_States_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("CountryUnited_States_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_CountryUnited_States_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_CountryUnited_States_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_CountryUnited_States_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_CountryUnited_States_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_CountryGermany --------------------- #### 34
def Crop_GOLD_CountryGermany(): 
    # ---------------- Dash1
    with Image.open("CountryGermany_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_CountryGermany_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("CountryGermany_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_CountryGermany_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_CountryGermany_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_CountryGermany_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_CountryGermany_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_CountryUnited_Kingdom --------------------- #### 35
def Crop_GOLD_CountryUnited_Kingdom(): 
    # ---------------- Dash1
    with Image.open("CountryUnited_Kingdom_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_CountryUnited_Kingdom_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("CountryUnited_Kingdom_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_CountryUnited_Kingdom_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_CountryUnited_Kingdom_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_CountryUnited_Kingdom_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_CountryUnited_Kingdom_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_CountrySouth_Korea --------------------- #### 36
def Crop_GOLD_CountrySouth_Korea(): 
    # ---------------- Dash1
    with Image.open("CountrySouth_Korea_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_CountrySouth_Korea_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("CountrySouth_Korea_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_CountrySouth_Korea_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_CountrySouth_Korea_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_CountrySouth_Korea_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_CountrySouth_Korea_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_CountryJapan --------------------- #### 37
def Crop_GOLD_CountryJapan(): 
    # ---------------- Dash1
    with Image.open("CountryJapan_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_CountryJapan_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("CountryJapan_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_CountryJapan_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_CountryJapan_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_CountryJapan_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_CountryJapan_dash2_bleu.png', quality=95) # BLEU

#### ----------------  GOLD_CountryFrance --------------------- #### 38
def Crop_GOLD_CountryFrance(): 
    # ---------------- Dash1
    with Image.open("CountryFrance_dash1.png") as im:
        im_Crop = im.crop( (0, 230, 1450, 2800) )
        im_Crop.save('crop_CountryFrance_dash1.png', quality=95)
        
    # ---------------- Dash2
    with Image.open("CountryFrance_dash2.png") as im2: 
        im_Crop2 = im2.crop((0, 250, 1450, 1320))
        im_Crop2.save('crop_CountryFrance_dash2_jaune.png', quality=95) # JAUNE
   
        im_Crop3 = im2.crop((0, 1320, 1450, 2020))
        im_Crop3.save('crop_CountryFrance_dash2_violet.png', quality=95) # VIOLET
         
        im_Crop4 = im2.crop((0, 2020, 1450, 2770))
        im_Crop4.save('crop_CountryFrance_dash2_vert.png', quality=95) # VERT
         
        im_Crop5 = im2.crop((0, 2770, 1450, 3910))
        im_Crop5.save('crop_CountryFrance_dash2_bleu.png', quality=95) # BLEU
        

############################################################




#---------------------- GET DOWNLOADABLE LINK -------------------- #######################################################
def DownloadLink(data_b64, doc_name):
    
    new_fileneme = f"{doc_name}_{PutDate}_dnl.docx"
    href = f'<a href="data:file/docx;base64,{data_b64}" Download="{new_fileneme}"> Download the doc file </a>'
    st.markdown(href, unsafe_allow_html=True)
    
### --------------- OPEN THE FILE TO LINK ----------------- ###

def fileOpen(filename):
        with open(filename, 'rb') as file:
            return base64.b64encode(file.read()).decode('UTF-8')

    
##########################################################################################################################


### -------------------- WORD DOCUMENTS -------------------- ###

#### ---------------- GOLD_CodirWorldwide --------------------- #### #1

def GOLD_CodirWorldwide_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_worldwide_dash1.png', width=Inches(6.5))

    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_worldwide_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_worldwide_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_worldwide_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_worldwide_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_worldwide_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_worldwide_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_worldwide_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_worldwide_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_CodirWorldwide.docx')
    document.save('GOLD_CodirWorldwide_{}_dnl.docx'.format(PutDate))
    

#### ---------------- GOLD_ZoneNorth_Asia --------------------- #### #2
def GOLD_ZoneNorth_Asia_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_ZoneNorth_Asia_dash1.png', width=Inches(6.5))
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_ZoneNorth_Asia_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_ZoneNorth_Asia_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_ZoneNorth_Asia_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_ZoneNorth_Asia_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_ZoneNorth_Asia.docx')
    document.save('GOLD_ZoneNorth_Asia_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_ZoneNorth_America --------------------- #### #3 
def ZoneNorth_America_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_ZoneNorth_America_dash1.png', width=Inches(6.5))
   
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_ZoneNorth_America_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_ZoneNorth_America_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_ZoneNorth_America_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_ZoneNorth_America_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_ZoneNorth_America.docx')
    document.save('GOLD_ZoneNorth_America_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_ZoneTravel_Retail --------------------- #### #4
def GOLD_ZoneTravel_Retail_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_ZoneTravel_Retail_dash1.png', width=Inches(6.5))
   
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_ZoneTravel_Retail_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_ZoneTravel_Retail_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_ZoneTravel_Retail_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_ZoneTravel_Retail_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_ZoneTravel_Retail.docx')
    document.save('GOLD_ZoneTravel_Retail_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_ZoneWestern_Europe --------------------- #### #5
def GOLD_ZoneWestern_Europe_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_ZoneWestern_Europe_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_ZoneWestern_Europe_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_ZoneWestern_Europe_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_ZoneWestern_Europe_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_ZoneWestern_Europe_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_ZoneWestern_Europe.docx')
    document.save('GOLD_ZoneWestern_Europe_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_ZoneEastern_Europe --------------------- #### #6
def GOLD_ZoneEastern_Europe_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_ZoneEastern_Europe_dash1.png', width=Inches(6.5))
   
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_ZoneEastern_Europe_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_ZoneEastern_Europe_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_ZoneEastern_Europe_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_ZoneEastern_Europe_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_ZoneEastern_Europe.docx')
    document.save('GOLD_ZoneEastern_Europe_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_ZoneSAPMENA --------------------- #### #7
def GOLD_ZoneSAPMENA_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_ZoneSAPMENA_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_ZoneSAPMENA_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_ZoneSAPMENA_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_ZoneSAPMENA_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_ZoneSAPMENA_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_ZoneSAPMENA.docx')
    document.save('GOLD_ZoneSAPMENA_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_ZoneLatin_America --------------------- #### #8
def GOLD_ZoneLatin_America_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_ZoneLatin_America_dash1.png', width=Inches(6.5))
   
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_ZoneLatin_America_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_ZoneLatin_America_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_ZoneLatin_America_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_ZoneLatin_America_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_ZoneLatin_America.docx')
    document.save('GOLD_ZoneLatin_America_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_ZoneSSA --------------------- #### #9
def GOLD_ZoneSSA_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_ZoneSSA_dash1.png', width=Inches(6.5))
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_ZoneSSA_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_ZoneSSA_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_ZoneSSA_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_ZoneSSA_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_ZoneSSA.docx')
    document.save('GOLD_ZoneSSA_{}_dnl.docx'.format(PutDate))


#### ----------------  GOLD_BrandKiehls --------------------- #### #10
def GOLD_BrandKiehls_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandKiehls_dash1.png', width=Inches(6.5))

    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandKiehls_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandKiehls_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandKiehls_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandKiehls_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandKiehls_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandKiehls_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandKiehls_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandKiehls_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandKiehls.docx')
    document.save('GOLD_BrandKiehls_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_BrandYves_Saint_Laurent --------------------- #### #11
def GOLD_BrandYves_Saint_Lauren_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandYves_Saint_Laurent_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandYves_Saint_Laurent_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandYves_Saint_Laurent_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandYves_Saint_Laurent_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandYves_Saint_Laurent_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandYves_Saint_Laurent_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandYves_Saint_Laurent_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandYves_Saint_Laurent_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandYves_Saint_Laurent_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandYves_Saint_Laurent.docx')
    document.save('GOLD_BrandYves_Saint_Laurent_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_BrandArmani --------------------- #### #12
def GOLD_BrandArmani_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandArmani_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandArmani_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandArmani_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandArmani_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandArmani_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandArmani_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandArmani_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandArmani_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandArmani_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandArmani.docx')
    document.save('GOLD_BrandArmani_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_BrandH_Rubinstein --------------------- #### #13
def GOLD_BrandH_Rubinstein_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandH_Rubinstein_dash1.png', width=Inches(6.5))
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandH_Rubinstein_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandH_Rubinstein_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandH_Rubinstein_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandH_Rubinstein_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandH_Rubinstein_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandH_Rubinstein_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandH_Rubinstein_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandH_Rubinstein_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandH_Rubinstein.docx')
    document.save('GOLD_BrandH_Rubinstein_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandBiotherm --------------------- #### #14
def GOLD_BrandBiotherm_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandBiotherm_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandBiotherm_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandBiotherm_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandBiotherm_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandBiotherm_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandBiotherm_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandBiotherm_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandBiotherm_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandBiotherm_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandBiotherm.docx')
    document.save('GOLD_BrandBiotherm_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandIT_Cosmetics --------------------- #### #15
def GOLD_BrandIT_Cosmetics_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandIT_Cosmetics_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandIT_Cosmetics_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandIT_Cosmetics_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandIT_Cosmetics_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandIT_Cosmetics_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandIT_Cosmetics_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandIT_Cosmetics_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandIT_Cosmetics_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandIT_Cosmetics_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandIT_Cosmetics.docx')
    document.save('GOLD_BrandIT_Cosmetics_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandUrban_Decay --------------------- #### 16
def GOLD_BrandUrban_Decay_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandUrban_Decay_dash1.png', width=Inches(6.5))

    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandUrban_Decay_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandUrban_Decay_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandUrban_Decay_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandUrban_Decay_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandUrban_Decay_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandUrban_Decay_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandUrban_Decay_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandUrban_Decay_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandUrban_Decay.docx')
    document.save('GOLD_BrandUrban_Decay_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandShu_Uemura --------------------- #### 17
def GOLD_BrandShu_Uemura_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandShu_Uemura_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandShu_Uemura_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandShu_Uemura_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandShu_Uemura_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandShu_Uemura_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandShu_Uemura_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandShu_Uemura_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandShu_Uemura_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandShu_Uemura_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandShu_Uemura.docx')
    document.save('GOLD_BrandShu_Uemura_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandRalph_Lauren --------------------- #### 18
def GOLD_BrandRalph_Lauren_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandRalph_Lauren_dash1.png', width=Inches(6.5))

    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandRalph_Lauren_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandRalph_Lauren_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandRalph_Lauren_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandRalph_Lauren_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandRalph_Lauren_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandRalph_Lauren_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandRalph_Lauren_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandRalph_Lauren_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandRalph_Lauren.docx')
    document.save('GOLD_BrandRalph_Lauren_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandYue_Sai --------------------- #### 19
def GOLD_BrandYue_Sai_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandYue_Sai_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandYue_Sai_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandYue_Sai_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandYue_Sai_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandRalph_Lauren_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandYue_Sai_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandYue_Sai_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandYue_Sai_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandYue_Sai_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandYue_Sai.docx')
    document.save('GOLD_BrandYue_Sai_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandValentino --------------------- #### 20
def GOLD_BrandValentino_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandValentino_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandValentino_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandValentino_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandValentino_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandValentino_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandValentino_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandValentino_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandValentino_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandValentino_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandValentino.docx')
    document.save('GOLD_BrandValentino_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandViktor_and_Rolf --------------------- #### 21
def GOLD_BrandViktor_and_Rolf_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandViktor_and_Rolf_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandViktor_and_Rolf_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandViktor_and_Rolf_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandViktor_and_Rolf_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandViktor_and_Rolf_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandViktor_and_Rolf_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandViktor_and_Rolf_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandViktor_and_Rolf_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandViktor_and_Rolf_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandViktor_and_Rolf.docx')
    document.save('GOLD_BrandViktor_and_Rolf_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandAtelier_Cologne --------------------- #### 22
def GOLD_BrandAtelier_Cologne_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandAtelier_Cologne_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandAtelier_Cologne_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandAtelier_Cologne_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandAtelier_Cologne_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandAtelier_Cologne_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandAtelier_Cologne_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandAtelier_Cologne_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandAtelier_Cologne_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandAtelier_Cologne_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandAtelier_Cologne.docx')
    document.save('GOLD_BrandAtelier_Cologne_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandMaison_Margiela --------------------- #### 23
def GOLD_BrandMaison_Margiela_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandMaison_Margiela_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandMaison_Margiela_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandMaison_Margiela_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandMaison_Margiela_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandMaison_Margiela_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandMaison_Margiela_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandMaison_Margiela_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandMaison_Margiela_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandMaison_Margiela_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandMaison_Margiela.docx')
    document.save('GOLD_BrandMaison_Margiela_{}_dnl.docx'.format(PutDate))

#### ---------------- GOLD_BrandMugler --------------------- #### 24
def GOLD_BrandMugler_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandMugler_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandMugler_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandMugler_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandMugler_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandMugler_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandMugler_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandMugler_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandMugler_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandMugler_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandMugler.docx')
    document.save('GOLD_BrandMugler_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_BrandCacharel --------------------- #### 25
def GOLD_BrandCacharel_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandCacharel_dash1.png', width=Inches(6.5))

    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandCacharel_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandCacharel_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandCacharel_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandCacharel_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandCacharel_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandCacharel_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandCacharel_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandCacharel_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandCacharel.docx')
    document.save('GOLD_BrandCacharel_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_BrandDiesel --------------------- #### 26
def GOLD_BrandDiesel_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandDiesel_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandDiesel_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandDiesel_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandDiesel_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandDiesel_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandDiesel_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandDiesel_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandDiesel_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandDiesel_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandDiesel.docx')
    document.save('GOLD_BrandDiesel_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_BrandAzzaro --------------------- #### 27
def GOLD_BrandAzzaro_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandAzzaro_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandAzzaro_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandAzzaro_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandAzzaro_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandAzzaro_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandAzzaro_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandAzzaro_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandAzzaro_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandAzzaro_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandAzzaro.docx')
    document.save('GOLD_BrandAzzaro_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_BrandPrada --------------------- #### 28
def GOLD_BrandPrada_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandPrada_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandPrada_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandPrada_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandPrada_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandPrada_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandPrada_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandPrada_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandPrada_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandPrada_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandPrada.docx')
    document.save('GOLD_BrandPrada_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_BrandHouse_99 --------------------- #### 29
def GOLD_BrandHouse_99_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandHouse_99_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandHouse_99_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandHouse_99_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandHouse_99_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandHouse_99_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandHouse_99_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandHouse_99_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandHouse_99_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandHouse_99_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandHouse_99.docx')
    document.save('GOLD_BrandHouse_99_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_BrandPaloma_Picasso --------------------- #### 30
def GOLD_BrandPaloma_Picasso_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandPaloma_Picasso_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandPaloma_Picasso_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandPaloma_Picasso_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandPaloma_Picasso_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandPaloma_Picasso_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandPaloma_Picasso_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandPaloma_Picasso_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandPaloma_Picasso_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandPaloma_Picasso_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandPaloma_Picasso.docx')
    document.save('GOLD_BrandPaloma_Picasso_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_BrandProenza_Schouler --------------------- #### 31
def GOLD_BrandProenza_Schouler_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_BrandProenza_Schouler_dash1.png', width=Inches(6.5))
    
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_BrandProenza_Schouler_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandProenza_Schouler_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_BrandProenza_Schouler_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_BrandProenza_Schouler_dash2_bleu.png', width=Cm(26.5))

    document.add_picture('crop_BrandProenza_Schouler_dash3_jaune.png', width=Cm(28))
    document.add_picture('crop_BrandProenza_Schouler_dash3_violet.png', width=Cm(29))
    document.add_picture('crop_BrandProenza_Schouler_dash3_vert.png', width=Cm(29))
    document.add_picture('crop_BrandProenza_Schouler_dash3_bleu.png', width=Cm(26.5))
    document.save('GOLD_BrandProenza_Schouler.docx')
    document.save('GOLD_BrandProenza_Schouler_{}_dnl.docx'.format(PutDate))


#### ----------------  GOLD_CountryChina --------------------- #### 32
def GOLD_CountryChina_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_CountryChina_dash1.png', width=Inches(6.5))
    
   
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_CountryChina_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_CountryChina_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_CountryChina_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_CountryChina_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_CountryChina.docx')
    document.save('GOLD_CountryChina_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_CountryUnited_States --------------------- #### 33
def GOLD_CountryUnited_States_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_CountryUnited_States_dash1.png', width=Inches(6.5))
    
   
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_CountryUnited_States_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_CountryUnited_States_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_CountryUnited_States_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_CountryUnited_States_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_CountryUnited_States.docx')
    document.save('GOLD_CountryUnited_States_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_CountryGermany --------------------- #### 34
def GOLD_CountryGermany_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_CountryGermany_dash1.png', width=Inches(6.5))
    
   
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_CountryGermany_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_CountryGermany_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_CountryGermany_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_CountryGermany_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_CountryGermany.docx')
    document.save('GOLD_CountryGermany_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_CountryUnited_Kingdom --------------------- #### 35
def GOLD_CountryUnited_Kingdom_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_CountryUnited_Kingdom_dash1.png', width=Inches(6.5))
    
   
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_CountryUnited_Kingdom_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_CountryUnited_Kingdom_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_CountryUnited_Kingdom_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_CountryUnited_Kingdom_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_CountryUnited_Kingdom.docx')
    document.save('GOLD_CountryUnited_Kingdom_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_CountrySouth_Korea --------------------- #### 36
def GOLD_CountrySouth_Korea_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_CountrySouth_Korea_dash1.png', width=Inches(6.5))
    

    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_CountrySouth_Korea_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_CountrySouth_Korea_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_CountrySouth_Korea_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_CountrySouth_Korea_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_CountrySouth_Korea.docx')
    document.save('GOLD_CountrySouth_Korea_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_CountryJapan --------------------- #### 37
def GOLD_CountryJapan_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_CountryJapan_dash1.png', width=Inches(6.5))
    
  
  
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_CountryJapan_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_CountryJapan_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_CountryJapan_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_CountryJapan_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_CountryJapan.docx')
    document.save('GOLD_CountryJapan_{}_dnl.docx'.format(PutDate))

#### ----------------  GOLD_CountryFrance --------------------- #### 38
def GOLD_CountryFrance_DOC():
    document = Document("pg_de_G.docx")
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
    document.add_picture('crop_CountryFrance_dash1.png', width=Inches(6.5))
    
   
   
    
    ### Orientation page
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height
    ### Orientation page to Landscape
    
    document.add_picture('crop_CountryFrance_dash2_jaune.png', width=Cm(28))
    document.add_picture('crop_CountryFrance_dash2_violet.png', width=Cm(29))
    document.add_picture('crop_CountryFrance_dash2_vert.png', width=Cm(29))
    document.add_picture('crop_CountryFrance_dash2_bleu.png', width=Cm(26.5))
    document.save('GOLD_CountryFrance.docx')
    document.save('GOLD_CountryFrance_{}_dnl.docx'.format(PutDate))

    
    
    
#############################################################################






####################### ZIP LINK DOWNLOAD ###################################
def Download_ZIP(data_b64, doc_name):
    
    new_fileneme = f"{doc_name}_{PutDate}.zip"
    href = f'<a href="data:file/zip;base64,{data_b64}" Download="{new_fileneme}"> Download All in Zip File </a>'
    st.markdown(href, unsafe_allow_html=True)
    
### --------------- Zip All Downloaded Docx ----------------- ###

def ZipAllDocx():
       with zipfile.ZipFile('final.zip', 'w') as zipF:
        for x in os.listdir():
            if "dnl.docx" in x:
                zipF.write(x, compress_type=zipfile.ZIP_DEFLATED)
                
### ---------------- Open Zip File ---------------------------- ###
  
def ZipOpen(filename):
    with open(filename, 'rb') as zipFile:
        return base64.b64encode(zipFile.read()).decode('UTF-8')

#####################################################################################



## --------------- APP APPARENCE OPTIONS ---------------- ####
st.set_page_config(
    page_title="GOLD Exporter L'Oréal",
    page_icon="📤",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://www.extremelycoolapp.com/help',
        'Report a bug': "https://www.extremelycoolapp.com/bug",
        'About': "# This is a header. This is an *extremely* cool app!"
     }
)


st.sidebar.image("LOREAL_LUXE_DIGITAL_BLACK.png", width=(250)) 
c1, c2, c3 = st.columns([1,2,1])
with c2:
    st.markdown("<h2 style='text-align: center; background-color: black; color:gold'>GOLD Exporter</h2>", unsafe_allow_html=True)

st.sidebar.markdown("<h4 style='text-align: left; color:slateblue'>DOMO AUTHENTIFICATION</h4>", unsafe_allow_html=True)
PutDate = st.sidebar.text_input("Select the date to put in documents")
USERNAME = st.sidebar.text_input("Login DOMO")
PASSWORD = st.sidebar.text_input("Password DOMO", type="password")
if st.sidebar.button('Save temporarily'):
    Username = USERNAME
    Password = PASSWORD


    


    
# ------------------ GETING SCREENSHOTS --------------------- #

def getscreenshot_of_url(url): 
    usr = (USERNAME)
    pwd = (PASSWORD)
    
    #url = ('https://loreal.domo.com/page/1178706243')
    
    ####### URLS ################
    
    #### ---------------- GOLD_CodirWorldwide --------------------- #### #1
    GOLD_CodirWorldwide_dash1 = 'https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    GOLD_CodirWorldwide_dash2 = 'https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    GOLD_CodirWorldwide_dash3 = 'https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Brand%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    
    #### ---------------- GOLD_ZoneNorth Asia --------------------- #### #2
    GOLD_ZoneNorth_Asia_dash1 = 'https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22North%20Asia%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    GOLD_ZoneNorth_Asia_dash2 = 'https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22North%20Asia%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    
    #### ----------------  GOLD_ZoneNorth_America --------------------- #### #3
    GOLD_ZoneNorth_America_dash1 = 'https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22North%20America%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    GOLD_ZoneNorth_America_dash2 = 'https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22North%20America%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    
    #### ----------------  GOLD_ZoneTravel_Retail --------------------- #### #4
    GOLD_ZoneTravel_Retail_dash1 = 'https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Travel%20Retail%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    GOLD_ZoneTravel_Retail_dash2 = 'https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Travel%20Retail%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'

    
    
    #### ----------------  GOLD_ZoneWestern_Europe --------------------- #### #5
    GOLD_ZoneWestern_Europe_dash1 = 'https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Western%20Europe%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    GOLD_ZoneWestern_Europe_dash2 = 'https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Western%20Europe%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    
    #### ----------------  GOLD_ZoneEastern_Europee --------------------- #### #6
    GOLD_ZoneEastern_Europe_dash1 = 'https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Eastern%20Europe%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    GOLD_ZoneEastern_Europe_dash2 = 'https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Eastern%20Europe%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
     
    #### ----------------  GOLD_ZoneSAPMENA --------------------- #### #7
    GOLD_ZoneSAPMENA_dash1 = 'https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22SAPMENA%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    GOLD_ZoneSAPMENA_dash2 = 'https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22SAPMENA%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    #### ----------------  GOLD_ZoneLatin_America --------------------- #### #8
    GOLD_ZoneLatin_America_dash1 = 'https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Latin%20America%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    GOLD_ZoneLatin_America_dash2 = 'https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Latin%20America%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    #### ----------------  GOLD_ZoneSSA --------------------- #### #9
    GOLD_ZoneSSA_dash1 = 'https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22SSA%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    GOLD_ZoneSSA_dash2 = 'https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Zone%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22SSA%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D'
    
    #### ----------------  GOLD_BrandKiehls --------------------- #### #10
    GOLD_BrandKiehls_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Kiehl%27s%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandKiehls_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Kiehl%27s%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandKiehls_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Kiehl%27s%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive%20Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_BrandYves_Saint_Lauren --------------------- #### #11
    GOLD_BrandYves_Saint_Laurent_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Yves Saint Laurent%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandYves_Saint_Laurent_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Yves Saint Laurent%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandYves_Saint_Laurent_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Yves Saint Laurent%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_BrandArmani --------------------- #### #12
    GOLD_BrandArmani_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Armani%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandArmani_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Armani%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandArmani_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Armani%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    

    #### ----------------  GOLD_BrandH_Rubinstein --------------------- #### #13
    GOLD_BrandH_Rubinstein_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22H. Rubinstein%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandH_Rubinstein_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22H. Rubinstein%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandH_Rubinstein_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22H. Rubinstein%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ---------------- GOLD_BrandBiotherm --------------------- #### #14
    GOLD_BrandBiotherm_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Biotherm%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandBiotherm_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Biotherm%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandBiotherm_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Biotherm%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ---------------- GOLD_BrandIT_Cosmetics --------------------- #### #15
    GOLD_BrandIT_Cosmetics_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22IT Cosmetics%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandIT_Cosmetics_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22IT Cosmetics%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandIT_Cosmetics_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22IT Cosmetics%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ---------------- GOLD_BrandUrban_Decay --------------------- #### 16
    GOLD_BrandUrban_Decay_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Urban Decay%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandUrban_Decay_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Urban Decay%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandUrban_Decay_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Urban Decay%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ---------------- GOLD_BrandShu_Uemura --------------------- #### 17
    GOLD_BrandShu_Uemura_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Shu Uemura%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandShu_Uemura_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Shu Uemura%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandShu_Uemura_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Shu Uemura%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    #### ---------------- GOLD_BrandRalph_Lauren --------------------- #### 18
    GOLD_BrandRalph_Lauren_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Ralph Lauren%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandRalph_Lauren_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Ralph Lauren%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandRalph_Lauren_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Ralph Lauren%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ---------------- GOLD_BrandYue_Sai --------------------- #### 19
    GOLD_BrandYue_Sai_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Yue-Sai%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandYue_Sai_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Yue-Sai%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandYue_Sai_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Yue-Sai%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ---------------- GOLD_BrandValentino --------------------- #### 20
    GOLD_BrandValentino_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Valentino%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandValentino_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Valentino%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandValentino_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Valentino%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ---------------- GOLD_BrandViktor_and_Rolf --------------------- #### 21 ###########################################################
    GOLD_BrandViktor_and_Rolf_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Viktor & Rolf%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandViktor_and_Rolf_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Viktor & Rolf%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandViktor_and_Rolf_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Viktor & Rolf%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    #### ---------------- GOLD_BrandAtelier_Cologne --------------------- #### 22
    GOLD_BrandAtelier_Cologne_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Atelier Cologne%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandAtelier_Cologne_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Atelier%20Cologne%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandAtelier_Cologne_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Atelier%20Cologne%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive%20Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ---------------- GOLD_BrandMaison_Margiela --------------------- #### 23
    GOLD_BrandMaison_Margiela_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Maison Margiela%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandMaison_Margiela_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Maison Margiela%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandMaison_Margiela_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Maison Margiela%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ---------------- GOLD_BrandMugler --------------------- #### 24
    GOLD_BrandMugler_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Mugler%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandMugler_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Mugler%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandMugler_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Mugler%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_BrandCacharel --------------------- #### 25
    GOLD_BrandCacharel_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Cacharel%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandCacharel_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Cacharel%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandCacharel_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Cacharel%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_BrandDiesel --------------------- #### 26
    GOLD_BrandDiesel_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Diesel%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandDiesel_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Diesel%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandDiesel_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Diesel%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_BrandAzzaro --------------------- #### 27
    GOLD_BrandAzzaro_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Azzaro%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandAzzaro_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Azzaro%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandAzzaro_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Azzaro%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_BrandPrada --------------------- #### 28
    GOLD_BrandPrada_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Prada%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandPrada_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Prada%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandPrada_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Prada%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_BrandHouse_99 --------------------- #### 29
    GOLD_BrandHouse_99_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22House 99%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandHouse_99_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22House 99%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandHouse_99_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22House 99%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_BrandPaloma_Picasso --------------------- #### 30
    GOLD_BrandPaloma_Picasso_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Paloma Picasso%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandPaloma_Picasso_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Paloma Picasso%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandPaloma_Picasso_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Paloma Picasso%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  OLD_BrandProenza_Schouler --------------------- #### 31
    GOLD_BrandProenza_Schouler_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Proenza Schouler%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandProenza_Schouler_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Proenza Schouler%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Zone%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_BrandProenza_Schouler_dash3 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Brand%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Proenza Schouler%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Drive Country%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_CountryChina --------------------- #### 32
    GOLD_CountryChina_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22China%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_CountryChina_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22China%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Brand%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_CountryUnited_States --------------------- #### 33
    GOLD_CountryUnited_States_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22United States%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_CountryUnited_States_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22United States%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Brand%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_CountryGermany --------------------- #### 34
    GOLD_CountryGermany_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Germany%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_CountryGermany_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Germany%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Brand%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_CountryUnited_Kingdom --------------------- #### 35
    GOLD_CountryUnited_Kingdom_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22United Kingdom%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_CountryUnited_Kingdom_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22United Kingdom%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Brand%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_CountrySouth_Korea --------------------- #### 36
    GOLD_CountrySouth_Korea_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22South Korea%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_CountrySouth_Korea_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22South Korea%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Brand%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_CountryJapan --------------------- #### 37
    GOLD_CountryJapan_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Japan%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_CountryJapan_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Japan%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Brand%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    #### ----------------  GOLD_CountryFrance --------------------- #### 38
    GOLD_CountryFrance_dash1 = "https://loreal.domo.com/page/1650379218?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22France%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    GOLD_CountryFrance_dash2 = "https://loreal.domo.com/page/798311033?pfilters=%5B%7B%22column%22:%22Country%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22France%22%5D%7D,%7B%22column%22:%22Breakdown%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22IN%22,%22values%22:%5B%22Brand%22%5D%7D,%7B%22column%22:%22GOLD%20-%20YTD%22,%22dataSourceId%22:%22%22,%22dataType%22:%22string%22,%22operand%22:%22NOT_EQUALS%22,%22values%22:%5B%22%22%5D%7D%5D"
    
    
    
    list_dash = [GOLD_CodirWorldwide_dash1, GOLD_CodirWorldwide_dash2, GOLD_CodirWorldwide_dash3, #OK 
                 
                 GOLD_ZoneNorth_Asia_dash1, GOLD_ZoneNorth_Asia_dash2, #OK
                 
                GOLD_ZoneNorth_America_dash1, GOLD_ZoneNorth_America_dash2, #OK
                 
                GOLD_ZoneTravel_Retail_dash1, GOLD_ZoneTravel_Retail_dash2, #OK
                
                GOLD_ZoneWestern_Europe_dash1, GOLD_ZoneWestern_Europe_dash2, #OK
                
                GOLD_ZoneEastern_Europe_dash1, GOLD_ZoneEastern_Europe_dash2, #OK
                
                GOLD_ZoneSAPMENA_dash1, GOLD_ZoneSAPMENA_dash2, #OK
                
                GOLD_ZoneLatin_America_dash1, GOLD_ZoneLatin_America_dash2, #OK
                
                GOLD_ZoneSSA_dash1, GOLD_ZoneSSA_dash2, #OK
                
                GOLD_BrandKiehls_dash1, GOLD_BrandKiehls_dash2, GOLD_BrandKiehls_dash3, #OK
                
                GOLD_BrandYves_Saint_Laurent_dash1, GOLD_BrandYves_Saint_Laurent_dash2, GOLD_BrandYves_Saint_Laurent_dash3, #OK
                
                GOLD_BrandArmani_dash1, GOLD_BrandArmani_dash2, GOLD_BrandArmani_dash3, #OK
                
                GOLD_BrandH_Rubinstein_dash1, GOLD_BrandH_Rubinstein_dash2, GOLD_BrandH_Rubinstein_dash3, #OK
                
                GOLD_BrandBiotherm_dash1, GOLD_BrandBiotherm_dash2, GOLD_BrandBiotherm_dash3, #OK
                
                GOLD_BrandIT_Cosmetics_dash1, GOLD_BrandIT_Cosmetics_dash2, GOLD_BrandIT_Cosmetics_dash3, #OK
                
                GOLD_BrandUrban_Decay_dash1, GOLD_BrandUrban_Decay_dash2, GOLD_BrandUrban_Decay_dash3, #OK
                 
                GOLD_BrandShu_Uemura_dash1, GOLD_BrandShu_Uemura_dash2, GOLD_BrandShu_Uemura_dash3, #OK

                GOLD_BrandRalph_Lauren_dash1, GOLD_BrandRalph_Lauren_dash2, GOLD_BrandRalph_Lauren_dash3, #OK

                GOLD_BrandYue_Sai_dash1, GOLD_BrandYue_Sai_dash2, GOLD_BrandYue_Sai_dash3, #OK

                GOLD_BrandValentino_dash1, GOLD_BrandValentino_dash2, GOLD_BrandValentino_dash3, #OK

                GOLD_BrandViktor_and_Rolf_dash1, GOLD_BrandViktor_and_Rolf_dash2, GOLD_BrandViktor_and_Rolf_dash3, #OK

                GOLD_BrandAtelier_Cologne_dash1, GOLD_BrandAtelier_Cologne_dash2, GOLD_BrandAtelier_Cologne_dash3, #OK

                GOLD_BrandMaison_Margiela_dash1, GOLD_BrandMaison_Margiela_dash2, GOLD_BrandMaison_Margiela_dash3, #OK

                GOLD_BrandMugler_dash1, GOLD_BrandMugler_dash2, GOLD_BrandMugler_dash3, #OK

                GOLD_BrandCacharel_dash1, GOLD_BrandCacharel_dash2, GOLD_BrandCacharel_dash3, #OK

                GOLD_BrandDiesel_dash1, GOLD_BrandDiesel_dash2, GOLD_BrandDiesel_dash3, #OK

                GOLD_BrandAzzaro_dash1, GOLD_BrandAzzaro_dash2, GOLD_BrandAzzaro_dash3, #OK

                GOLD_BrandPrada_dash1, GOLD_BrandPrada_dash2, GOLD_BrandPrada_dash3, #OK
                
                GOLD_BrandHouse_99_dash1, GOLD_BrandHouse_99_dash2, GOLD_BrandHouse_99_dash3, #OK

                GOLD_BrandPaloma_Picasso_dash1, GOLD_BrandPaloma_Picasso_dash2, GOLD_BrandPaloma_Picasso_dash3, #OK

                GOLD_BrandProenza_Schouler_dash1, GOLD_BrandProenza_Schouler_dash2, GOLD_BrandProenza_Schouler_dash3, #OK

                GOLD_CountryChina_dash1, GOLD_CountryChina_dash2, #OK
                
                GOLD_CountryUnited_States_dash1, GOLD_CountryUnited_States_dash2, #OK

                GOLD_CountryGermany_dash1, GOLD_CountryGermany_dash2, #OK

                GOLD_CountryUnited_Kingdom_dash1, GOLD_CountryUnited_Kingdom_dash2, #OK

                GOLD_CountrySouth_Korea_dash1, GOLD_CountrySouth_Korea_dash2, #OK

                GOLD_CountryJapan_dash1, GOLD_CountryJapan_dash2, #OK

                GOLD_CountryFrance_dash1, GOLD_CountryFrance_dash2] #OK
    
##############
    
    chrome_options = Options()
    chrome_options.headless = False
    chrome_options.add_argument("--incognito")
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--start-maximized")
    
    
    

    
    #'./chromedriver', options=chrome_options
    driver = webdriver.Chrome("chromedrive.exe", options=chrome_options)
    driver.get(GOLD_CodirWorldwide_dash1)
    
    
    
    #####################
    login_button = driver.find_element_by_class_name('db-button')
    login_button.click()
    driver.implicitly_wait(50)
    
    username_box = driver.find_element_by_name('loginfmt')
    username_box.send_keys(usr)
    driver.implicitly_wait(50)
    
    suivant = driver.find_element_by_id('idSIButton9')
    suivant.click()
    driver.implicitly_wait(50)
    
    password_box = driver.find_element_by_name('Password') 
    password_box.send_keys(pwd)
    driver.implicitly_wait(50)
    
    connexion = driver.find_element_by_id('submitButton') #submitButton 
    connexion.click()
    driver.implicitly_wait(50)
    
    validate_auth = driver.find_element_by_class_name('table')
    validate_auth.click()
    #driver.implicitly_wait(20)
    time.sleep(65)

    closeBuzz = driver.find_element_by_id('BuzzAnchor')
    closeBuzz.click()
    driver.implicitly_wait(50)

    driver.implicitly_wait(70) # gives an implicit wait for 20 seconds

    element = driver.find_element_by_xpath('//*[@id="pg-layout"]/div') # The same everywhere in domo dashboards 
    

    driver.implicitly_wait(50)
    width = 3000
    height = element.size['height'] + 900
    driver.set_window_size(width,height)
    #driver.maximize_window()
    
    
    
    
    

    for u in list_dash:
#### ---------------- GOLD_CodirWorldwide --------------------- 1

        if u == (GOLD_CodirWorldwide_dash1):
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            img = driver.save_screenshot("worldwide_dash1.png")
            driver.implicitly_wait(30)

        elif u == (GOLD_CodirWorldwide_dash2):
            driver.get(GOLD_CodirWorldwide_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("worldwide_dash2.png")
            driver.implicitly_wait(30)

        elif u == (GOLD_CodirWorldwide_dash3):
            driver.get(GOLD_CodirWorldwide_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("worldwide_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_CodirWorldwide()
            # ----------- Doc Word ---------
            GOLD_CodirWorldwide_DOC()
            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_CodirWorldwide_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_CodirWorldwide.docx')
                DownloadLink(my_doc, 'GOLD_CodirWorldwide')
            #with col4:
                
                
#### ---------------- ZoneNorth_Asia --------------------- 2
        elif u == (GOLD_ZoneNorth_Asia_dash1):
            driver.get(GOLD_ZoneNorth_Asia_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneNorth_Asia_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_ZoneNorth_Asia_dash2):
            driver.get(GOLD_ZoneNorth_Asia_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneNorth_Asia_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_ZoneNorth_Asia()
            # ----------- Doc Word ---------
            GOLD_ZoneNorth_Asia_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_ZoneNorth_Asia_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_ZoneNorth_Asia.docx')
                DownloadLink(my_doc, 'GOLD_ZoneNorth_Asia')
            
 #### ---------------- GOLD_ZoneNorth_America --------------------- 3
           
        elif u == (GOLD_ZoneNorth_America_dash1):
            driver.get(GOLD_ZoneNorth_America_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneNorth_America_dash1.png")
            driver.implicitly_wait(30)
            
        elif u ==(GOLD_ZoneNorth_America_dash2):
            driver.get(GOLD_ZoneNorth_America_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneNorth_America_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_ZoneNorth_America()
            # ----------- Doc Word ---------
            ZoneNorth_America_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_ZoneNorth_America_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_ZoneNorth_America.docx')
                DownloadLink(my_doc, 'GOLD_ZoneNorth_America')
            
#### ---------------- GOLD_ZoneTravel_Retail --------------------- 4
        elif u == (GOLD_ZoneTravel_Retail_dash1):
            driver.get(GOLD_ZoneTravel_Retail_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneTravel_Retail_dash1.png")
            driver.implicitly_wait(30)
        
        elif u == (GOLD_ZoneTravel_Retail_dash2):
            driver.get(GOLD_ZoneTravel_Retail_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneTravel_Retail_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_ZoneTravel_Retail()
            # ----------- Doc Word ---------
            GOLD_ZoneTravel_Retail_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_ZoneTravel_Retail_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_ZoneTravel_Retail.docx')
                DownloadLink(my_doc, 'GOLD_ZoneTravel_Retail')
        
#### ---------------- GOLD_ZoneWestern_Europe --------------------- 5       
        elif u == (GOLD_ZoneWestern_Europe_dash1):
            driver.get(GOLD_ZoneWestern_Europe_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneWestern_Europe_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_ZoneWestern_Europe_dash2):
            driver.get(GOLD_ZoneWestern_Europe_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneWestern_Europe_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_ZoneWestern_Europe()
            # ----------- Doc Word ---------
            GOLD_ZoneWestern_Europe_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_ZoneWestern_Europe_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_ZoneWestern_Europe.docx')
                DownloadLink(my_doc, 'GOLD_ZoneWestern_Europe')
            
#### ---------------- GOLD_ZoneEastern_Europe --------------------- 6 
            
        elif u == (GOLD_ZoneEastern_Europe_dash1):
            driver.get(GOLD_ZoneEastern_Europe_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneEastern_Europe_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_ZoneEastern_Europe_dash2):
            driver.get(GOLD_ZoneEastern_Europe_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneEastern_Europe_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_ZoneEastern_Europe()
            # ----------- Doc Word ---------
            GOLD_ZoneEastern_Europe_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_ZoneEastern_Europe_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_ZoneEastern_Europe.docx')
                DownloadLink(my_doc, 'GOLD_ZoneEastern_Europe')
            
#### ---------------- GOLD_ZoneSAPMENA --------------------- 7
            
        elif u == (GOLD_ZoneSAPMENA_dash1):
            driver.get(GOLD_ZoneSAPMENA_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneSAPMENA_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_ZoneSAPMENA_dash2):
            driver.get(GOLD_ZoneSAPMENA_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneSAPMENA_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_ZoneSAPMENA()
            # ----------- Doc Word ---------
            GOLD_ZoneSAPMENA_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_ZoneSAPMENA_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_ZoneSAPMENA.docx')
                DownloadLink(my_doc, 'GOLD_ZoneSAPMENA')
            
#### ---------------- GOLD_ZoneLatin_America --------------------- 8
            
        elif u == (GOLD_ZoneLatin_America_dash1):
            driver.get(GOLD_ZoneLatin_America_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneLatin_America_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_ZoneLatin_America_dash2):
            driver.get(GOLD_ZoneLatin_America_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneLatin_America_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_ZoneLatin_America()
            # ----------- Doc Word ---------
            GOLD_ZoneLatin_America_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_ZoneLatin_America_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_ZoneLatin_America.docx')
                DownloadLink(my_doc, 'GOLD_ZoneLatin_America')
            
#### ---------------- GOLD_ZoneSSA --------------------- 9 
        elif u == (GOLD_ZoneSSA_dash1):
            driver.get(GOLD_ZoneSSA_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneSSA_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_ZoneSSA_dash2):
            driver.get(GOLD_ZoneSSA_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("ZoneSSA_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_ZoneSSA()
            # ----------- Doc Word ---------
            GOLD_ZoneSSA_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_ZoneSSA_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_ZoneSSA.docx')
                DownloadLink(my_doc, 'GOLD_ZoneSSA')
            
#### ---------------- GOLD_BrandKiehls --------------------- 10   
        elif u == (GOLD_BrandKiehls_dash1):
            driver.get(GOLD_BrandKiehls_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandKiehls_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandKiehls_dash2):
            driver.get(GOLD_BrandKiehls_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandKiehls_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandKiehls_dash3):
            driver.get(GOLD_BrandKiehls_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandKiehls_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandKiehls()
            # ----------- Doc Word ---------
            GOLD_BrandKiehls_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandKiehls_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandKiehls.docx')
                DownloadLink(my_doc, 'GOLD_BrandKiehls')
            
#### ---------------- GOLD_BrandYves_Saint_Laurent --------------------- 11
        
        elif u == (GOLD_BrandYves_Saint_Laurent_dash1):
            driver.get(GOLD_BrandYves_Saint_Laurent_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandYves_Saint_Laurent_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandYves_Saint_Laurent_dash2):
            driver.get(GOLD_BrandYves_Saint_Laurent_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandYves_Saint_Laurent_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandYves_Saint_Laurent_dash3):
            driver.get(GOLD_BrandYves_Saint_Laurent_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandYves_Saint_Laurent_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandYves_Saint_Lauren()
            # ----------- Doc Word ---------
            GOLD_BrandYves_Saint_Lauren_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandYves_Saint_Laurent_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandYves_Saint_Laurent.docx')
                DownloadLink(my_doc, 'GOLD_BrandYves_Saint_Laurent')
            
#### ---------------- GOLD_BrandArmani --------------------- 12
            
        elif u == (GOLD_BrandArmani_dash1):
            driver.get(GOLD_BrandArmani_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandArmani_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandArmani_dash2):
            driver.get(GOLD_BrandArmani_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandArmani_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandArmani_dash3):
            driver.get(GOLD_BrandArmani_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandArmani_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandArmani()
            # ----------- Doc Word ---------
            GOLD_BrandArmani_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandArmani_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandArmani.docx')
                DownloadLink(my_doc, 'GOLD_BrandArmani')
            
#### ---------------- GOLD_BrandH_Rubinstein --------------------- 13
            
        elif u == (GOLD_BrandH_Rubinstein_dash1):
            driver.get(GOLD_BrandH_Rubinstein_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandH_Rubinstein_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandH_Rubinstein_dash2):
            driver.get(GOLD_BrandH_Rubinstein_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandH_Rubinstein_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandH_Rubinstein_dash3):
            driver.get(GOLD_BrandH_Rubinstein_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandH_Rubinstein_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandH_Rubinstein()
            # ----------- Doc Word ---------
            GOLD_BrandH_Rubinstein_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandH_Rubinstein_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandH_Rubinstein.docx')
                DownloadLink(my_doc, 'GOLD_BrandH_Rubinstein')
            
#### ---------------- GOLD_BrandBiotherm --------------------- 14
            
        elif u == (GOLD_BrandBiotherm_dash1):
            driver.get(GOLD_BrandBiotherm_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandBiotherm_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandBiotherm_dash2):
            driver.get(GOLD_BrandBiotherm_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandBiotherm_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandBiotherm_dash3):
            driver.get(GOLD_BrandBiotherm_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandBiotherm_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandBiotherm()
            # ----------- Doc Word ---------
            GOLD_BrandBiotherm_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandBiotherm_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandBiotherm.docx')
                DownloadLink(my_doc, 'GOLD_BrandBiotherm')
            
#### ---------------- GOLD_BrandIT_Cosmetics --------------------- 15
            
        elif u == (GOLD_BrandIT_Cosmetics_dash1):
            driver.get(GOLD_BrandIT_Cosmetics_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandIT_Cosmetics_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandIT_Cosmetics_dash2):
            driver.get(GOLD_BrandIT_Cosmetics_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandIT_Cosmetics_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandIT_Cosmetics_dash3):
            driver.get(GOLD_BrandIT_Cosmetics_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandIT_Cosmetics_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandIT_Cosmetics()
            # ----------- Doc Word ---------
            GOLD_BrandIT_Cosmetics_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandIT_Cosmetics_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandIT_Cosmetics.docx')
                DownloadLink(my_doc, 'GOLD_BrandIT_Cosmetics')
            
#### ---------------- GOLD_BrandUrban_Decay --------------------- 16
            
        elif u == (GOLD_BrandUrban_Decay_dash1):
            driver.get(GOLD_BrandUrban_Decay_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandUrban_Decay_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandUrban_Decay_dash2):
            driver.get(GOLD_BrandUrban_Decay_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandUrban_Decay_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandUrban_Decay_dash3):
            driver.get(GOLD_BrandUrban_Decay_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandUrban_Decay_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandUrban_Decay()
            # ----------- Doc Word ---------
            GOLD_BrandUrban_Decay_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandUrban_Decay_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandUrban_Decay.docx')
                DownloadLink(my_doc, 'GOLD_BrandUrban_Decay')
            
#### ---------------- GOLD_BrandShu_Uemura --------------------- 17
            
        elif u == (GOLD_BrandShu_Uemura_dash1):
            driver.get(GOLD_BrandShu_Uemura_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandShu_Uemura_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandShu_Uemura_dash2):
            driver.get(GOLD_BrandShu_Uemura_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandShu_Uemura_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandShu_Uemura_dash3):
            driver.get(GOLD_BrandShu_Uemura_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandShu_Uemura_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandShu_Uemura()
            # ----------- Doc Word ---------
            GOLD_BrandShu_Uemura_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandShu_Uemura_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandShu_Uemura.docx')
                DownloadLink(my_doc, 'GOLD_BrandShu_Uemura')
            
#### ---------------- GOLD_BrandRalph_Lauren --------------------- 18
            
        elif u == (GOLD_BrandRalph_Lauren_dash1):
            driver.get(GOLD_BrandRalph_Lauren_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandRalph_Lauren_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandRalph_Lauren_dash2):
            driver.get(GOLD_BrandRalph_Lauren_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandRalph_Lauren_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandRalph_Lauren_dash3):
            driver.get(GOLD_BrandRalph_Lauren_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandRalph_Lauren_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandRalph_Lauren()
            # ----------- Doc Word ---------
            GOLD_BrandRalph_Lauren_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandRalph_Lauren_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandRalph_Lauren.docx')
                DownloadLink(my_doc, 'GOLD_BrandRalph_Lauren')
            
#### ---------------- GOLD_BrandYue_Sai --------------------- 19
            
        elif u == (GOLD_BrandYue_Sai_dash1):
            driver.get(GOLD_BrandYue_Sai_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandYue_Sai_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandYue_Sai_dash2):
            driver.get(GOLD_BrandYue_Sai_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandYue_Sai_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandYue_Sai_dash3):
            driver.get(GOLD_BrandYue_Sai_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandYue_Sai_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandYue_Sai()
            # ----------- Doc Word ---------
            GOLD_BrandYue_Sai_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandYue_Sai_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandYue_Sai.docx')
                DownloadLink(my_doc, 'GOLD_BrandYue_Sai')
            
#### ---------------- GOLD_BrandValentino --------------------- 20
            
        elif u == (GOLD_BrandValentino_dash1):
            driver.get(GOLD_BrandValentino_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandValentino_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandValentino_dash2):
            driver.get(GOLD_BrandValentino_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandValentino_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandValentino_dash3):
            driver.get(GOLD_BrandValentino_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandValentino_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandValentino()
            # ----------- Doc Word ---------
            GOLD_BrandValentino_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandValentino_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandValentino.docx')
                DownloadLink(my_doc, 'GOLD_BrandValentino')
            
#### ---------------- GOLD_BrandViktor_and_Rolf --------------------- 21
            
        elif u == (GOLD_BrandViktor_and_Rolf_dash1):
            driver.get(GOLD_BrandViktor_and_Rolf_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandViktor_and_Rolf_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandViktor_and_Rolf_dash2):
            driver.get(GOLD_BrandViktor_and_Rolf_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandViktor_and_Rolf_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandViktor_and_Rolf_dash3):
            driver.get(GOLD_BrandViktor_and_Rolf_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandViktor_and_Rolf_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandViktor_and_Rolf()
            # ----------- Doc Word ---------
            GOLD_BrandViktor_and_Rolf_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandViktor_and_Rolf_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandViktor_and_Rolf.docx')
                DownloadLink(my_doc, 'GOLD_BrandViktor_and_Rolf')
            
#### ---------------- GOLD_BrandAtelier_Cologne --------------------- 22
            
        elif u == (GOLD_BrandAtelier_Cologne_dash1):
            driver.get(GOLD_BrandAtelier_Cologne_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandAtelier_Cologne_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandAtelier_Cologne_dash2):
            driver.get(GOLD_BrandAtelier_Cologne_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandAtelier_Cologne_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandAtelier_Cologne_dash3):
            driver.get(GOLD_BrandAtelier_Cologne_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandAtelier_Cologne_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandAtelier_Cologne()
            # ----------- Doc Word ---------
            GOLD_BrandAtelier_Cologne_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandAtelier_Cologne_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandAtelier_Cologne.docx')
                DownloadLink(my_doc, 'GOLD_BrandAtelier_Cologne')
            
#### ---------------- GOLD_BrandMaison_Margiela --------------------- 23
            
        elif u == (GOLD_BrandMaison_Margiela_dash1):
            driver.get(GOLD_BrandMaison_Margiela_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandMaison_Margiela_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandMaison_Margiela_dash2):
            driver.get(GOLD_BrandMaison_Margiela_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandMaison_Margiela_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandMaison_Margiela_dash3):
            driver.get(GOLD_BrandMaison_Margiela_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandMaison_Margiela_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandMaison_Margiela()
            # ----------- Doc Word ---------
            GOLD_BrandMaison_Margiela_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandMaison_Margiela_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandMaison_Margiela.docx')
                DownloadLink(my_doc, 'GOLD_BrandMaison_Margiela')
            
#### ---------------- GOLD_BrandMugler --------------------- 24
            
        elif u == (GOLD_BrandMugler_dash1):
            driver.get(GOLD_BrandMugler_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandMugler_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandMugler_dash2):
            driver.get(GOLD_BrandMugler_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandMugler_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandMugler_dash3):
            driver.get(GOLD_BrandMugler_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandMugler_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandMugler()
            # ----------- Doc Word ---------
            GOLD_BrandMugler_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandMugler_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandMugler.docx')
                DownloadLink(my_doc, 'GOLD_BrandMugler')
            
#### ---------------- GOLD_BrandCacharel --------------------- 25
            
        elif u == (GOLD_BrandCacharel_dash1):
            driver.get(GOLD_BrandCacharel_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandCacharel_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandCacharel_dash2):
            driver.get(GOLD_BrandCacharel_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandCacharel_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandCacharel_dash3):
            driver.get(GOLD_BrandCacharel_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandCacharel_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandCacharel()
            # ----------- Doc Word ---------
            GOLD_BrandCacharel_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandCacharel_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandCacharel.docx')
                DownloadLink(my_doc, 'GOLD_BrandCacharel')
            
#### ---------------- GOLD_BrandDiesel --------------------- 26
            
        elif u == (GOLD_BrandDiesel_dash1):
            driver.get(GOLD_BrandDiesel_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandDiesel_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandDiesel_dash2):
            driver.get(GOLD_BrandDiesel_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandDiesel_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandDiesel_dash3):
            driver.get(GOLD_BrandDiesel_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandDiesel_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandDiesel()
            # ----------- Doc Word ---------
            GOLD_BrandDiesel_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandDiesel_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandDiesel.docx')
                DownloadLink(my_doc, 'GOLD_BrandDiesel')
            
#### ---------------- GOLD_BrandAzzaro --------------------- 27
            
        elif u == (GOLD_BrandAzzaro_dash1):
            driver.get(GOLD_BrandAzzaro_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandAzzaro_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandAzzaro_dash2):
            driver.get(GOLD_BrandAzzaro_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandAzzaro_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandAzzaro_dash3):
            driver.get(GOLD_BrandAzzaro_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandAzzaro_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandAzzaro()
            # ----------- Doc Word ---------
            GOLD_BrandAzzaro_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandAzzaro_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandAzzaro.docx')
                DownloadLink(my_doc, 'GOLD_BrandAzzaro')
            
#### ---------------- GOLD_BrandPrada --------------------- 28
            
        elif u == (GOLD_BrandPrada_dash1):
            driver.get(GOLD_BrandPrada_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandPrada_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandPrada_dash2):
            driver.get(GOLD_BrandPrada_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandPrada_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandPrada_dash3):
            driver.get(GOLD_BrandPrada_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandPrada_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandPrada()
            # ----------- Doc Word ---------
            GOLD_BrandPrada_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandPrada_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandPrada.docx')
                DownloadLink(my_doc, 'GOLD_BrandPrada')
            
#### ---------------- GOLD_BrandHouse_99 --------------------- 29
            
        elif u == (GOLD_BrandHouse_99_dash1):
            driver.get(GOLD_BrandHouse_99_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandHouse_99_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandHouse_99_dash2):
            driver.get(GOLD_BrandHouse_99_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandHouse_99_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandHouse_99_dash3):
            driver.get(GOLD_BrandHouse_99_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandHouse_99_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandHouse_99()
            # ----------- Doc Word ---------
            GOLD_BrandHouse_99_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandHouse_99_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandHouse_99.docx')
                DownloadLink(my_doc, 'GOLD_BrandHouse_99')
            
#### ---------------- GOLD_BrandPaloma_Picasso --------------------- 30
            
        elif u == (GOLD_BrandPaloma_Picasso_dash1):
            driver.get(GOLD_BrandPaloma_Picasso_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandPaloma_Picasso_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandPaloma_Picasso_dash2):
            driver.get(GOLD_BrandPaloma_Picasso_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandPaloma_Picasso_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandPaloma_Picasso_dash3):
            driver.get(GOLD_BrandPaloma_Picasso_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandPaloma_Picasso_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandPaloma_Picasso()
            # ----------- Doc Word ---------
            GOLD_BrandPaloma_Picasso_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandPaloma_Picasso_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandPaloma_Picasso.docx')
                DownloadLink(my_doc, 'GOLD_BrandPaloma_Picasso')
            
#### ---------------- GOLD_BrandProenza_Schouler --------------------- 31
            
        elif u == (GOLD_BrandProenza_Schouler_dash1):
            driver.get(GOLD_BrandProenza_Schouler_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandProenza_Schouler_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandProenza_Schouler_dash2):
            driver.get(GOLD_BrandProenza_Schouler_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandProenza_Schouler_dash2.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_BrandProenza_Schouler_dash3):
            driver.get(GOLD_BrandProenza_Schouler_dash3)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("BrandProenza_Schouler_dash3.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_BrandProenza_Schouler()
            # ----------- Doc Word ---------
            GOLD_BrandProenza_Schouler_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_BrandProenza_Schouler_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_BrandProenza_Schouler.docx')
                DownloadLink(my_doc, 'GOLD_BrandProenza_Schouler')
            
#### ---------------- GOLD_CountryChina --------------------- 32
            
        elif u == (GOLD_CountryChina_dash1):
            driver.get(GOLD_CountryChina_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryChina_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_CountryChina_dash2):
            driver.get(GOLD_CountryChina_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryChina_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_CountryChina()
            # ----------- Doc Word ---------
            GOLD_CountryChina_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_CountryChina_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_CountryChina.docx')
                DownloadLink(my_doc, 'GOLD_CountryChina')
            
#### ---------------- GOLD_CountryUnited_States --------------------- 33
            
        elif u == (GOLD_CountryUnited_States_dash1):
            driver.get(GOLD_CountryUnited_States_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryUnited_States_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_CountryUnited_States_dash2):
            driver.get(GOLD_CountryUnited_States_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryUnited_States_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_CountryUnited_States()
            # ----------- Doc Word ---------
            GOLD_CountryUnited_States_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_CountryUnited_States_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_CountryUnited_States.docx')
                DownloadLink(my_doc, 'GOLD_CountryUnited_States')
            
#### ---------------- GOLD_CountryGermany --------------------- 34
            
        elif u == (GOLD_CountryGermany_dash1):
            driver.get(GOLD_CountryGermany_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryGermany_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_CountryGermany_dash2):
            driver.get(GOLD_CountryGermany_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryGermany_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_CountryGermany()
            # ----------- Doc Word ---------
            GOLD_CountryGermany_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_CountryGermany_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_CountryGermany.docx')
                DownloadLink(my_doc, 'GOLD_CountryGermany')
            
#### ---------------- GOLD_CountryUnited_Kingdom --------------------- 35
            
        elif u == (GOLD_CountryUnited_Kingdom_dash1):
            driver.get(GOLD_CountryUnited_Kingdom_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryUnited_Kingdom_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_CountryUnited_Kingdom_dash2):
            driver.get(GOLD_CountryUnited_Kingdom_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryUnited_Kingdom_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_CountryUnited_Kingdom()
            # ----------- Doc Word ---------
            GOLD_CountryUnited_Kingdom_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_CountryUnited_Kingdom_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_CountryUnited_Kingdom.docx')
                DownloadLink(my_doc, 'GOLD_CountryUnited_Kingdom')
            
#### ---------------- GOLD_CountrySouth_Korea --------------------- 36
            
        elif u == (GOLD_CountrySouth_Korea_dash1):
            driver.get(GOLD_CountrySouth_Korea_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountrySouth_Korea_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_CountrySouth_Korea_dash2):
            driver.get(GOLD_CountrySouth_Korea_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountrySouth_Korea_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_CountrySouth_Korea()
            # ----------- Doc Word ---------
            GOLD_CountrySouth_Korea_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_CountrySouth_Korea_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_CountrySouth_Korea.docx')
                DownloadLink(my_doc, 'GOLD_CountrySouth_Korea')
            
#### ---------------- GOLD_CountryJapan --------------------- 37
            
        elif u == (GOLD_CountryJapan_dash1):
            driver.get(GOLD_CountryJapan_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryJapan_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_CountryJapan_dash2):
            driver.get(GOLD_CountryJapan_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryJapan_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_CountryJapan()
            # ----------- Doc Word ---------
            GOLD_CountryJapan_DOC()

            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_CountryJapan_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_CountryJapan.docx')
                DownloadLink(my_doc, 'GOLD_CountryJapan')
            
#### ---------------- GOLD_CountryFrance --------------------- 38
            
        elif u == (GOLD_CountryFrance_dash1):
            driver.get(GOLD_CountryFrance_dash1)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryFrance_dash1.png")
            driver.implicitly_wait(30)
            
        elif u == (GOLD_CountryFrance_dash2):
            driver.get(GOLD_CountryFrance_dash2)
            driver.execute_script("document.body.style.zoom='90%'")
            time.sleep(90)
            
            img = driver.save_screenshot("CountryFrance_dash2.png")
            driver.implicitly_wait(30)
            # ----------Croping ------------
            Crop_GOLD_CountryFrance()
            # ----------- Doc Word ---------
            GOLD_CountryFrance_DOC()
            ###### Results in columns --------------------
            col1, col2, col3, col4 = st.columns([3,2,1,2])
            with col1:
                st.write("GOLD_CountryFrance_{}.docx".format(PutDate))
            with col2:
                my_doc = fileOpen('GOLD_CountryFrance.docx')
                DownloadLink(my_doc, 'GOLD_CountryFrance')
                
            with col4:
                ZipAllDocx()
                my_zip_doc = ZipOpen('final.zip')
                Download_ZIP(my_zip_doc, 'GOLD_Exporter_ZIP')


            driver.quit()

                   
################################################################

restrictions = ('![^]"#$#%&()*+,/:;<=>?')


                
def main():
    
    st.markdown(hide_menu, unsafe_allow_html = True)

    Nav_tab = ["Login", "SignUp"]
    st.sidebar.markdown("<h4 style='text-align: left; color:slateblue'> LOGIN TO THE APP</h4>", unsafe_allow_html=True)
    choice = st.sidebar.selectbox("Login/SignUp",Nav_tab)
    
    if choice == "Login": 
        st.markdown("<h2 style='text-align: center'></h2>", unsafe_allow_html=True)
            
       
        uname = st.sidebar.text_input("Username")

        ### Restrictions unmae
        if restrictions in uname: 
            st.error("Special characters not allowed !")
        elif len(uname) > 15:
            st.error("The Username must contain no more than 15 characters")

        passwd = st.sidebar.text_input("Password", type='password')

        ### Restrictions passwd
        if len(passwd) > 20:
            st.error("The Password must contains no more than 20 characters")
        else:
        #### End passwd restrictions
            if st.sidebar.checkbox('Login'):

                create_usertable()
                hashed_pswd = make_hashes(passwd)
                result = login_user(uname, check_hashes(passwd,hashed_pswd))
                if result:
                    st.success("Hello {}! You are connected... Make sure to disconnect before leaving the App.".format(uname))

                    #st.write('<style>div.row-widget.stRadio> div{flex-direction:row;}<style>', unsafe_allow_html=True)
                    task = st.selectbox("", ["Documentation 📁", "Exporter 📤", 
                                             "Customizable Crop 📷", "Customizable Link 🔁" ,"Profiles Manager 👥"])

                    if task == "Documentation 📁":
                        ## Guide d'utilisation
                        st.markdown("<h2 style='text-align: center'> 📋 Please read this before using the App </h2>", unsafe_allow_html=True)
                        st.info("""Before exporting dashboards,
                        please follow the following steps:""")
                        st.info("""

                        1. Go to the Exporter section

                        2. Select the date in the sidebar to put in yours documents. *The default date is the current date.

                        3. Put your DOMO ID's in the sidebar

                        4. Click on Save temporarily.""")
                        st.warning("""SECURITY: Once the export is finished, make sure to delete your DOMO ID's before to disconnect""")

                    elif task == "Exporter 📤":

                        st.markdown("<h2 style='text-align: center'></h2>", unsafe_allow_html=True)

                        if st.button("Start Exporting"):
                            st.markdown("<h3 style='text-align: left'>Documents</h3>", unsafe_allow_html=True)
                            result = getscreenshot_of_url("url")
                            # We will put here dynamic link
                            st.markdown("<h3 style='text-align: center'>Visualization</h3>", unsafe_allow_html=True)
                            st.image('crop_worldwide_dash2_jaune.png')


                    elif task == "Customizable Crop 📷":
                        st.markdown("<h2 style='text-align: center'> Customizable Crop</h2>", unsafe_allow_html=True)
                        st.text("Available in Version 2")


                    elif task == "Customizable Link 🔁":
                        st.markdown("<h2 style='text-align: center'> Customizable Link</h2>", unsafe_allow_html=True)
                        st.text("Available in Version 2")


                    elif task == "Profiles Manager 👥":
                        st.markdown("<h2 style='text-align: center'> Profiles Manager</h2>", unsafe_allow_html=True)

                        user_result = view_all_users()
                        clean_db = pd.DataFrame(user_result, columns=["Username", "Password"])
                        st.dataframe(clean_db)

                else:
                    st.error("Incorrect username/password")

            else:
                col1, col2, col3 = st.columns([1,2,3.5])
                with col2:
                    st.image("LOREAL_LUXE_DIGITAL_BLACK.png", width=(600))
                st.markdown("<h5 style='text-align: center'>Hello, please go to the Login section to log in</h5>", unsafe_allow_html=True)
            
                
    elif choice == "SignUp":
        st.markdown("<h2 style='text-align: center'> Create New Account</h2>", unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([2,3,2])
        with col2:
            new_user = st.text_input("Username")
            new_password = st.text_input("Password", type="password")
            
            if restrictions in new_user: 
                st.error("Special characters not allowed !")
                
            elif len(new_user) > 15:
                st.error("The Username must contain no more than 15 characters")
                
            elif len(new_password) > 20:
                st.error("The Password must contains no more than 20 characters")
                
            else:
                if st.button("Create Account"):
                    create_usertable()
                    add_userdata(new_user, make_hashes(new_password))
                    st.success("Votre compte a été créé")
                    st.info("Go back to the Login Section to log in...")

    
if __name__ == '__main__':
    main()


    
    
    
    
    
    
    
    
    
    
    
    
    
